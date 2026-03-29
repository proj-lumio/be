"""
Generate mock contract .docx files for prodotto3 through prodotto20.
Each prodotto has a unique SaaS product, category, and client naming theme.
"""

import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE_DIR = os.path.join(os.path.dirname(__file__), "..", "mock_data")

# ── Product definitions ──
# (folder, product_name, product_description, category, canone, soglia_or_metric, date, extra_details)

PRODUCTS = [
    {
        "folder": "prodotto3",
        "name": "VaultCRM",
        "category": "CRM e gestione clienti",
        "desc": "una piattaforma CRM in modalità SaaS finalizzata alla gestione centralizzata delle relazioni con i clienti, al tracciamento delle opportunità commerciali, alla segmentazione della clientela, alla gestione delle pipeline di vendita e alla generazione di report commerciali",
        "canone": 1500,
        "soglia": "10 utenti",
        "extra_unit": 25,
        "retention_days": 730,
        "date": "15/04/2025",
        "sla_target": "99,7%",
        "sla_min": "98,5%",
        "credito_pct": 8,
        "resp_limit_pct": 25,
        "clients": [
            "Eins", "Dva", "Tre_cz", "Quatre", "Cinco", "Sechs", "Sette_de",
            "Oito", "Dokuz", "Tien", "Jedenaście", "Tolv", "Dreizehn",
            "Quatorze", "Femten", "Sedici_ro", "Dix_sept", "Eighteen",
            "Devetnaest", "Dwadziescia"
        ],
    },
    {
        "folder": "prodotto4",
        "name": "ShieldNet",
        "category": "Cybersecurity e monitoraggio",
        "desc": "una piattaforma SaaS di monitoraggio continuo della sicurezza informatica, finalizzata al rilevamento di minacce, all'analisi delle vulnerabilità, alla gestione degli incidenti di sicurezza, al monitoraggio del traffico di rete e alla generazione di alert in tempo reale",
        "canone": 3000,
        "soglia": "500 endpoint",
        "extra_unit": 4,
        "retention_days": 365,
        "date": "01/02/2025",
        "sla_target": "99,9%",
        "sla_min": "99,0%",
        "credito_pct": 10,
        "resp_limit_pct": 30,
        "clients": [
            "Rosso", "Azul", "Gelb", "Vert", "Arancione", "Lila", "Bianco",
            "Siyah", "Dorado", "Srebrny", "Brun", "Roz", "Turkuaz",
            "Groen", "Karmazyn", "Celeste", "Bordo", "Ivory", "Ocra", "Grigio"
        ],
    },
    {
        "folder": "prodotto5",
        "name": "InvoiceFlow",
        "category": "Fatturazione elettronica",
        "desc": "una piattaforma SaaS per la gestione della fatturazione elettronica, l'emissione e la ricezione di fatture in formato XML conforme al Sistema di Interscambio (SDI), la conservazione digitale a norma, la riconciliazione automatica dei pagamenti e la generazione di report fiscali",
        "canone": 800,
        "soglia": "500 fatture/mese",
        "extra_unit": 0.15,
        "retention_days": 3650,
        "date": "10/03/2025",
        "sla_target": "99,5%",
        "sla_min": "98,0%",
        "credito_pct": 10,
        "resp_limit_pct": 20,
        "clients": [
            "Lupo", "Katze", "Ours", "Aguila", "Cervo", "Tavsan", "Delfin",
            "Falke", "Leone", "Papuga", "Bjorn", "Colibri", "Tigre_pt",
            "Vrabec", "Elefante", "Gepard", "Lachs", "Pavone", "Gorila", "Renna"
        ],
    },
    {
        "folder": "prodotto6",
        "name": "PeopleHub",
        "category": "Gestione risorse umane",
        "desc": "una piattaforma SaaS per la gestione integrata delle risorse umane, comprensiva di funzionalità di anagrafica dipendenti, gestione presenze e assenze, elaborazione cedolini, pianificazione ferie, valutazione delle performance, onboarding digitale e reportistica HR",
        "canone": 1200,
        "soglia": "50 dipendenti",
        "extra_unit": 8,
        "retention_days": 1825,
        "date": "22/05/2025",
        "sla_target": "99,5%",
        "sla_min": "98,0%",
        "credito_pct": 7,
        "resp_limit_pct": 20,
        "clients": [
            "Rosa_fi", "Tulipan", "Orchidee", "Girasol", "Lavanda", "Jasmin",
            "Dahlia", "Lilien", "Narciso", "Krokus", "Violeta", "Camelia",
            "Peonia", "Freesia", "Magnolia", "Azalea", "Begonia", "Zinnia",
            "Gardenia", "Anemone"
        ],
    },
    {
        "folder": "prodotto7",
        "name": "TrackLine",
        "category": "Logistica e supply chain",
        "desc": "una piattaforma SaaS per la gestione della logistica e della supply chain, finalizzata al tracciamento delle spedizioni in tempo reale, alla gestione del magazzino, all'ottimizzazione dei percorsi di consegna, alla pianificazione degli approvvigionamenti e alla generazione di KPI logistici",
        "canone": 2000,
        "soglia": "1.000 spedizioni/mese",
        "extra_unit": 0.80,
        "retention_days": 365,
        "date": "08/06/2025",
        "sla_target": "99,7%",
        "sla_min": "98,5%",
        "credito_pct": 10,
        "resp_limit_pct": 25,
        "clients": [
            "Wien", "Praha", "Lisboa", "Athen", "Bern", "Vilnius", "Riga",
            "Tallinn", "Zagreb", "Ankara", "Київ", "Minsk", "Skopje",
            "Tirana", "Beograd", "Chisinau", "Podgorica", "Reykjavik",
            "Valletta", "Nicosia"
        ],
    },
    {
        "folder": "prodotto8",
        "name": "InsightBI",
        "category": "Business intelligence e analytics",
        "desc": "una piattaforma SaaS di business intelligence finalizzata alla raccolta, trasformazione e visualizzazione dei dati aziendali, alla creazione di dashboard interattive, all'analisi predittiva tramite modelli statistici e di machine learning, e alla generazione automatica di report direzionali",
        "canone": 2500,
        "soglia": "5 data source",
        "extra_unit": 200,
        "retention_days": 730,
        "date": "18/01/2025",
        "sla_target": "99,5%",
        "sla_min": "98,0%",
        "credito_pct": 8,
        "resp_limit_pct": 20,
        "clients": [
            "Orion", "Vega", "Sirius", "Altair", "Rigel", "Polaris",
            "Antares", "Deneb", "Cassiopea", "Andromeda", "Perseo", "Lyra",
            "Centauri", "Betelgeuse", "Aldebaran", "Spica", "Arcturus",
            "Canopus", "Procione", "Capella"
        ],
    },
    {
        "folder": "prodotto9",
        "name": "DocVault",
        "category": "Gestione documentale",
        "desc": "una piattaforma SaaS per la gestione documentale aziendale, finalizzata all'archiviazione, indicizzazione, ricerca full-text, versionamento, condivisione sicura e conservazione digitale a norma di documenti aziendali in qualsiasi formato, con workflow approvativi configurabili",
        "canone": 900,
        "soglia": "10 GB di storage",
        "extra_unit": 5,
        "retention_days": 3650,
        "date": "05/07/2025",
        "sla_target": "99,5%",
        "sla_min": "98,0%",
        "credito_pct": 10,
        "resp_limit_pct": 20,
        "clients": [
            "Quercia", "Birke", "Cedro", "Abete", "Olmo", "Platano",
            "Cipresso", "Kastanie", "Salice", "Pino", "Acero", "Frassino",
            "Baobab", "Sequoia", "Tasso", "Ontano", "Noce", "Tiglio",
            "Eucalipto", "Palma"
        ],
    },
    {
        "folder": "prodotto10",
        "name": "SprintDesk",
        "category": "Project management",
        "desc": "una piattaforma SaaS di project management finalizzata alla pianificazione, esecuzione e monitoraggio di progetti aziendali, comprensiva di gestione task, diagrammi di Gantt, board Kanban, time tracking, allocazione risorse e reportistica di avanzamento progetto",
        "canone": 1000,
        "soglia": "15 utenti",
        "extra_unit": 12,
        "retention_days": 365,
        "date": "28/02/2025",
        "sla_target": "99,5%",
        "sla_min": "98,0%",
        "credito_pct": 8,
        "resp_limit_pct": 20,
        "clients": [
            "Danubio", "Tamigi", "Senna", "Reno", "Volga", "Nilo", "Tevere",
            "Tago", "Elba", "Oder", "Vistola", "Loira", "Rodano", "Moldava",
            "Drava", "Mosa", "Nemunas", "Daugava", "Arno", "Adige"
        ],
    },
    {
        "folder": "prodotto11",
        "name": "ShopEngine",
        "category": "E-commerce",
        "desc": "una piattaforma SaaS per la creazione e gestione di negozi online, comprensiva di catalogo prodotti, carrello, checkout, gestione ordini, integrazione con gateway di pagamento, gestione magazzino, promozioni, coupon e analytics delle vendite",
        "canone": 1800,
        "soglia": "1.000 ordini/mese",
        "extra_unit": 0.50,
        "retention_days": 730,
        "date": "14/04/2025",
        "sla_target": "99,7%",
        "sla_min": "98,5%",
        "credito_pct": 10,
        "resp_limit_pct": 25,
        "clients": [
            "Everest", "Bianco_mt", "Cervino", "Olimpo", "Etna", "Fuji",
            "Kilimanjaro", "Elbrus", "Aconcagua", "Denali", "Matterhorn",
            "Vesuvio", "Sinai", "Ararat", "Kenia_mt", "Dolomiti", "Carpazi",
            "Balcani_mt", "Pirenei", "Alpi"
        ],
    },
    {
        "folder": "prodotto12",
        "name": "SensIoT",
        "category": "IoT e monitoraggio industriale",
        "desc": "una piattaforma SaaS per il monitoraggio industriale IoT, finalizzata alla raccolta dati da sensori e dispositivi connessi, alla visualizzazione in tempo reale di parametri operativi, alla gestione di alert e soglie, alla manutenzione predittiva e alla generazione di report di efficienza impianto",
        "canone": 3500,
        "soglia": "100 dispositivi",
        "extra_unit": 15,
        "retention_days": 365,
        "date": "20/08/2025",
        "sla_target": "99,9%",
        "sla_min": "99,0%",
        "credito_pct": 10,
        "resp_limit_pct": 30,
        "clients": [
            "Rubino", "Saphir", "Smeraldo", "Diamante", "Topazio", "Opale",
            "Ametista", "Granato", "Turchese", "Giada", "Agata", "Quarzo",
            "Zircone", "Lapislazzuli", "Corallo", "Ambra", "Perla", "Onice",
            "Tanzanite", "Malachite"
        ],
    },
    {
        "folder": "prodotto13",
        "name": "LexComply",
        "category": "Legal tech e compliance",
        "desc": "una piattaforma SaaS per la gestione della compliance normativa e legale, finalizzata al monitoraggio degli adempimenti, alla gestione delle scadenze regolamentari, all'analisi automatica di testi normativi, alla generazione di checklist di conformità e alla produzione di report per audit interni ed esterni",
        "canone": 2200,
        "soglia": "3 entità giuridiche",
        "extra_unit": 400,
        "retention_days": 3650,
        "date": "03/09/2025",
        "sla_target": "99,5%",
        "sla_min": "98,0%",
        "credito_pct": 10,
        "resp_limit_pct": 20,
        "clients": [
            "Mistral", "Tramontana", "Scirocco", "Bora", "Levante", "Ponente",
            "Libeccio", "Grecale", "Zefiro", "Boreas", "Chinook", "Foehn",
            "Harmattan", "Monsone", "Passat", "Brezza", "Tifone", "Uragano",
            "Tornado", "Ciclone"
        ],
    },
    {
        "folder": "prodotto14",
        "name": "MediConnect",
        "category": "Telemedicina e sanità digitale",
        "desc": "una piattaforma SaaS per la telemedicina e la gestione sanitaria digitale, finalizzata alla prenotazione di visite online, alla gestione delle cartelle cliniche elettroniche, alla teleconsulenza medico-paziente, al monitoraggio remoto dei parametri vitali e alla generazione di referti digitali",
        "canone": 4000,
        "soglia": "20 medici",
        "extra_unit": 80,
        "retention_days": 3650,
        "date": "11/10/2025",
        "sla_target": "99,9%",
        "sla_min": "99,5%",
        "credito_pct": 10,
        "resp_limit_pct": 30,
        "clients": [
            "Aquila_bd", "Rondine", "Fenice", "Pellicano", "Cardellino",
            "Usignolo", "Merlo", "Gabbiano", "Allodola", "Airone",
            "Falcone", "Gufo", "Colibrì", "Cigno", "Pettirosso",
            "Tortora", "Cuculo", "Corvo", "Fringuello", "Passero"
        ],
    },
    {
        "folder": "prodotto15",
        "name": "LearnSphere",
        "category": "EdTech e formazione",
        "desc": "una piattaforma SaaS per la formazione aziendale e l'e-learning, finalizzata alla creazione e distribuzione di corsi online, alla gestione di percorsi formativi personalizzati, alla somministrazione di test di valutazione, al tracciamento delle competenze acquisite e alla certificazione digitale",
        "canone": 1100,
        "soglia": "100 learner",
        "extra_unit": 5,
        "retention_days": 730,
        "date": "25/11/2025",
        "sla_target": "99,5%",
        "sla_min": "98,0%",
        "credito_pct": 8,
        "resp_limit_pct": 20,
        "clients": [
            "Mela", "Kirsche", "Naranja", "Citron", "Fragola", "Ananas",
            "Mango", "Pesca", "Kiwi_fr", "Melograno", "Banana", "Lampone",
            "Myrtille", "Fico", "Papaya", "Cocco", "Litchi", "Dattero",
            "Ribes", "Mora"
        ],
    },
    {
        "folder": "prodotto16",
        "name": "PayGate",
        "category": "FinTech e pagamenti",
        "desc": "una piattaforma SaaS per la gestione dei pagamenti digitali, finalizzata all'elaborazione di transazioni online, alla riconciliazione automatica dei flussi finanziari, alla gestione di wallet aziendali, all'integrazione con circuiti di pagamento nazionali e internazionali e alla reportistica finanziaria in tempo reale",
        "canone": 2000,
        "soglia": "5.000 transazioni/mese",
        "extra_unit": 0.10,
        "retention_days": 3650,
        "date": "07/12/2025",
        "sla_target": "99,9%",
        "sla_min": "99,5%",
        "credito_pct": 10,
        "resp_limit_pct": 30,
        "clients": [
            "Tango", "Valzer", "Samba", "Flamenco", "Polka", "Mazurka",
            "Bolero", "Salsa", "Rumba", "Merengue", "Bachata", "Foxtrot",
            "Quickstep", "Jive", "Charleston", "Tarantella", "Breakdance",
            "Cumbia", "Reggaeton", "Swing"
        ],
    },
    {
        "folder": "prodotto17",
        "name": "CampaignForge",
        "category": "Marketing automation",
        "desc": "una piattaforma SaaS per l'automazione del marketing, finalizzata alla gestione di campagne email, alla segmentazione avanzata del pubblico, al lead scoring, alla creazione di landing page, al tracciamento delle conversioni, all'A/B testing e alla generazione di report di performance delle campagne",
        "canone": 1300,
        "soglia": "10.000 contatti",
        "extra_unit": 0.02,
        "retention_days": 365,
        "date": "19/01/2026",
        "sla_target": "99,5%",
        "sla_min": "98,0%",
        "credito_pct": 8,
        "resp_limit_pct": 20,
        "clients": [
            "Violino", "Klavier", "Guitarra", "Flauto", "Trompete",
            "Harfe", "Saxofon", "Tamburo", "Cello", "Oboe", "Fagotto",
            "Klarinette", "Tuba", "Mandolino", "Banjo", "Organo",
            "Fisarmonica", "Marimba", "Sitar", "Didgeridoo"
        ],
    },
    {
        "folder": "prodotto18",
        "name": "CloudNest",
        "category": "Cloud storage e backup",
        "desc": "una piattaforma SaaS per lo storage cloud e il backup aziendale, finalizzata all'archiviazione sicura dei dati, alla sincronizzazione multi-dispositivo, al backup automatico incrementale, al disaster recovery, alla gestione granulare dei permessi di accesso e alla crittografia end-to-end dei dati",
        "canone": 700,
        "soglia": "100 GB",
        "extra_unit": 3,
        "retention_days": 365,
        "date": "02/03/2026",
        "sla_target": "99,9%",
        "sla_min": "99,0%",
        "credito_pct": 10,
        "resp_limit_pct": 25,
        "clients": [
            "Solstizio", "Equinozio", "Aurora_mt", "Tramonto", "Nebbia",
            "Tempesta", "Brina", "Rugiada", "Grandine", "Arcobaleno",
            "Eclissi", "Crepuscolo", "Sereno", "Nuvoloso", "Foschia",
            "Calura", "Gelo", "Pioggia", "Neve", "Tuono"
        ],
    },
    {
        "folder": "prodotto19",
        "name": "GreenLens",
        "category": "Sostenibilità e reporting ESG",
        "desc": "una piattaforma SaaS per il monitoraggio e il reporting ESG (Environmental, Social, Governance), finalizzata alla raccolta dati sulle emissioni di CO₂, alla gestione degli indicatori di sostenibilità, alla generazione di report conformi agli standard GRI, ESRS e CSRD, e al tracciamento degli obiettivi di decarbonizzazione",
        "canone": 2800,
        "soglia": "1 entità giuridica",
        "extra_unit": 1500,
        "retention_days": 3650,
        "date": "16/02/2026",
        "sla_target": "99,5%",
        "sla_min": "98,0%",
        "credito_pct": 8,
        "resp_limit_pct": 20,
        "clients": [
            "Ferro", "Kupfer", "Argento_el", "Platino", "Titanio", "Cobalto",
            "Nichel", "Cromo", "Vanadio", "Tungsteno", "Zinco", "Stagno",
            "Bismuto", "Litio", "Silicio", "Boro", "Cesio", "Iridio",
            "Osmio", "Rodio"
        ],
    },
    {
        "folder": "prodotto20",
        "name": "HelpStream",
        "category": "Customer support e ticketing",
        "desc": "una piattaforma SaaS per la gestione del customer support, finalizzata alla ricezione e gestione di ticket multicanale (email, chat, telefono, social), alla creazione di knowledge base, alla gestione di SLA interni, al routing intelligente delle richieste, alla chatbot automation e alla reportistica di customer satisfaction",
        "canone": 1500,
        "soglia": "10 agenti",
        "extra_unit": 30,
        "retention_days": 365,
        "date": "30/01/2026",
        "sla_target": "99,7%",
        "sla_min": "98,5%",
        "credito_pct": 10,
        "resp_limit_pct": 25,
        "clients": [
            "Risotto", "Croissant", "Pretzel", "Paella", "Sushi",
            "Goulash", "Pierogi", "Baklava", "Moussaka", "Ceviche",
            "Borscht", "Tagine", "Kimchi", "Ramen", "Tiramisu",
            "Strudel", "Falafel", "Empanada", "Fondue", "Gyoza"
        ],
    },
]


def build_contract_text(product, client_name):
    """Build the full contract text for a given product and client."""
    p = product
    return f"""CONTRATTO DI FORNITURA DI SERVIZI SOFTWARE IN MODALITÀ SaaS, SUPPORTO E MANUTENZIONE
{p['name']}

tra

CompletlyFake S.r.l., con sede legale in Milano, P. IVA IT12345678901, in persona del legale rappresentante pro tempore, di seguito il «Fornitore»

e

{client_name}, con sede legale in Milano, in persona del legale rappresentante pro tempore, di seguito il «Cliente»

congiuntamente, le «Parti» e, singolarmente, la «Parte».

1. Premesse

1.1 Le premesse e gli allegati costituiscono parte integrante e sostanziale del presente Contratto.

1.2 Il Fornitore sviluppa e mette a disposizione {p['desc']}.

1.3 Il Cliente intende utilizzare {p['name']} nell'ambito delle proprie attività aziendali, secondo i termini, le condizioni e i limiti previsti dal presente Contratto.

1.4 Il presente Contratto è sottoscritto in data {p['date']} e acquista efficacia dalla medesima data.

2. Definizioni

a) «Servizio»: il servizio SaaS denominato {p['name']}.
b) «Piattaforma»: l'infrastruttura applicativa e tecnologica con cui il Servizio è erogato.
c) «Dati del Cliente»: i contenuti, le informazioni e i dati caricati o generati dal Cliente nell'ambito dell'utilizzo del Servizio.
d) «Utenti Autorizzati»: i soggetti autorizzati dal Cliente ad accedere al Servizio.
e) «Canone di Licenza»: il corrispettivo fisso trimestrale dovuto dal Cliente per il diritto di utilizzo del Servizio.
f) «Corrispettivo Variabile»: il corrispettivo dovuto in base ai volumi eccedenti la soglia inclusa.
g) «Ticket Critico»: segnalazione relativa a indisponibilità totale del Servizio.
h) «Ticket Medio»: segnalazione relativa a compromissione di funzionalità importanti del Servizio.
i) «Ticket Basso»: segnalazione relativa a indisponibilità parziale o degrado limitato del Servizio, in presenza di workaround disponibile.
j) «Uptime»: percentuale di disponibilità del Servizio misurata su base continuativa 24 ore su 24, 7 giorni su 7, nel periodo di riferimento.
k) «Credito di Servizio»: qualsiasi sconto, accredito o riduzione del corrispettivo riconosciuto al Cliente in relazione al mancato rispetto degli SLA.
l) «Forza Maggiore»: qualsiasi evento al di fuori del ragionevole controllo del Fornitore, inclusi, a titolo esemplificativo, outage gravi e generalizzati di provider infrastrutturali cloud, inclusa Microsoft Azure, interruzioni massicce di rete, eventi naturali, atti dell'autorità, cyber attacchi sistemici o altri eventi straordinari.

3. Oggetto del Contratto

3.1 Con il presente Contratto il Fornitore concede al Cliente un diritto non esclusivo, non trasferibile e non sublicenziabile di accesso e utilizzo del Servizio {p['name']} in modalità SaaS, per la durata del Contratto e nei limiti ivi previsti.

3.2 Il Servizio rientra nella categoria «{p['category']}» e consente al Cliente di fruire delle funzionalità descritte nelle premesse, nei limiti del profilo commerciale attivato.

3.3 Salvo quanto diversamente previsto, il presente Contratto include: a) accesso al Servizio in modalità SaaS; b) hosting dell'ambiente in Microsoft Azure — West Europe; c) supporto tecnico tramite sistema di ticketing; d) manutenzione correttiva ordinaria del software standard.

4. Caratteristiche e limiti del Servizio

4.1 Il Cliente prende atto che il Servizio effettua elaborazioni automatizzate e che i risultati prodotti dalla Piattaforma possono variare in funzione della qualità, completezza e correttezza dei dati forniti in input.

4.2 Il Fornitore non garantisce che le elaborazioni siano prive di errori, omissioni o ambiguità in ogni caso d'uso.

4.3 Il Cliente resta responsabile della verifica dei risultati prima del loro utilizzo in processi decisionali, amministrativi, contrattuali, contabili, legali, regolamentari, di compliance o di altra natura.

4.4 Le performance del Servizio possono variare in funzione, a titolo esemplificativo, del volume dei dati, della complessità delle configurazioni, della qualità degli input e delle condizioni dell'infrastruttura sottostante.

5. Corrispettivi, fatturazione e pagamenti

5.1 A fronte dei servizi oggetto del presente Contratto, il Cliente corrisponderà al Fornitore: a) un Canone di Licenza fisso trimestrale pari a Euro {p['canone']}, fatturato trimestralmente in via anticipata, con pagamento a 30 giorni data fattura; b) un Corrispettivo Variabile a consumo, fatturato mensilmente posticipato, con pagamento a 30 giorni data fattura, determinato ai sensi dell'articolo 5.4.

5.2 Tutti gli importi si intendono al netto di IVA e di ogni altro onere di legge.

5.3 In caso di ritardo nei pagamenti, saranno dovuti gli interessi moratori ai sensi del D. Lgs. 231/2002, fatta salva la facoltà del Fornitore di sospendere il Servizio previa comunicazione scritta ai sensi dell'articolo 11.

5.4 Il Canone di Licenza include una soglia base di {p['soglia']}. Per i volumi eccedenti tale soglia, sarà dovuto un Corrispettivo Variabile pari a Euro {p['extra_unit']} per unità aggiuntiva nel periodo di riferimento.

5.5 Qualsiasi Credito di Servizio, sconto o riduzione del corrispettivo eventualmente spettante al Cliente ai sensi del presente Contratto si applicherà esclusivamente al Canone di Licenza fisso e non potrà in alcun caso essere calcolato sul, né compensato con, il Corrispettivo Variabile a consumo.

6. Livelli di servizio — disponibilità

6.1 Il Fornitore misurerà la disponibilità del Servizio su base 24x7, con rilevazione trimestrale.

6.2 Il Fornitore si impegna a perseguire: a) un obiettivo di disponibilità pari al {p['sla_target']} su base trimestrale; b) una soglia minima di servizio pari al {p['sla_min']} su base trimestrale.

6.3 Qualora, nel trimestre di riferimento, l'Uptime del Servizio risulti inferiore al {p['sla_min']}, il Cliente avrà diritto esclusivamente a un Credito di Servizio pari al {p['credito_pct']}% del Canone di Licenza del trimestre di riferimento, da applicarsi sulla fattura successiva.

6.4 Il Credito di Servizio di cui al presente articolo costituisce l'unico rimedio economico spettante al Cliente in relazione al mancato raggiungimento dei livelli di disponibilità del Servizio.

6.5 Ai fini del calcolo dell'Uptime non saranno considerati i periodi di indisponibilità imputabili a: a) eventi di Forza Maggiore; b) indisponibilità o outage gravi di Microsoft Azure o di altri provider infrastrutturali essenziali non ragionevolmente controllabili dal Fornitore; c) manutenzione programmata comunicata con congruo preavviso; d) malfunzionamenti della rete, dei sistemi o delle integrazioni del Cliente o di terzi; e) uso improprio del Servizio da parte del Cliente.

7. Livelli di servizio — ticketing e presa in carico

7.1 Il supporto tecnico sarà erogato tramite sistema di ticketing.

7.2 I ticket saranno classificati come segue: a) Livello 1 / Critico: sistema completamente non disponibile; b) Livello 2 / Medio: funzionalità importanti compromesse; c) Livello 3 / Basso: alcune funzionalità non disponibili, con workaround disponibile.

7.3 I tempi target di presa in carico, misurati in ore di calendario 24x7, saranno i seguenti: a) Ticket Critici: entro 4 ore; b) Ticket Medi: entro 8 ore; c) Ticket Bassi: entro 1 giorno di calendario.

7.4 Qualora, in un determinato trimestre, il Fornitore non prenda in carico entro i tempi SLA almeno il 95% dei ticket aperti nel trimestre, il Cliente avrà diritto a un Credito di Servizio pari al 5% del Canone di Licenza del trimestre successivo.

7.5 Il Credito di Servizio di cui all'articolo 7.4 maturerà solo ove, nel trimestre di riferimento, risultino aperti almeno 2 (due) ticket complessivamente classificati come Critici e/o Medi.

7.6 Anche per il presente SLA, il Credito di Servizio costituisce l'unico rimedio economico spettante al Cliente.

7.7 Sono esclusi dal computo dei tempi di presa in carico i ticket impropri, duplicati, privi di informazioni minime sufficienti o riconducibili a cause non imputabili al Fornitore.

8. Tetto massimo dei crediti di servizio

8.1 I Crediti di Servizio eventualmente maturati dal Cliente ai sensi degli articoli 6 e 7 potranno essere cumulati tra loro.

8.2 In ogni caso, l'ammontare complessivo dei Crediti di Servizio riconoscibili al Cliente in relazione al medesimo trimestre non potrà eccedere il {p['credito_pct']}% del Canone di Licenza del trimestre di riferimento.

8.3 Resta in ogni caso esclusa qualunque applicazione dei Crediti di Servizio al Corrispettivo Variabile a consumo.

9. Supporto e manutenzione

9.1 Il Fornitore fornirà manutenzione correttiva e supporto tecnico standard relativi al Servizio.

9.2 Non rientrano nel perimetro standard: a) manutenzione evolutiva; b) personalizzazioni; c) nuove funzionalità; d) nuove integrazioni; e) attività di configurazione specialistica non espressamente incluse nell'offerta economica.

9.3 Tali attività, ove richieste dal Cliente, formeranno oggetto di separata quotazione.

10. Data retention e cancellazione

10.1 I Dati del Cliente, nonché i relativi output, saranno conservati dal Fornitore per un periodo massimo di {p['retention_days']} giorni decorrenti dalla data di elaborazione.

10.2 Decorso tale termine, i dati saranno cancellati in via irreversibile e definitiva dai sistemi del Fornitore, fatti salvi eventuali obblighi di legge inderogabili.

10.3 Il Cliente prende atto e accetta che, decorso il termine di retention di cui sopra, il Fornitore non sarà tenuto a rendere nuovamente disponibili dati o output.

11. Hosting, localizzazione dati e sospensione

11.1 Il Servizio sarà erogato su infrastruttura Microsoft Azure con localizzazione dei dati in West Europe.

11.2 Il Cliente prende atto e accetta che il Fornitore si avvale di Microsoft Azure quale provider infrastrutturale primario del Servizio.

11.3 Ove il Fornitore tratti dati personali per conto del Cliente, l'eventuale ricorso a Microsoft Azure o ad altri fornitori strettamente necessari all'erogazione del Servizio sarà disciplinato nel DPA, nei limiti della normativa applicabile.

11.4 In caso di mancato pagamento di una fattura entro il termine pattuito, il Fornitore potrà sospendere, in tutto o in parte, il Servizio decorso inutilmente un termine di 15 (quindici) giorni dal ricevimento, da parte del Cliente, di un sollecito scritto di pagamento inviato dal Fornitore.

12. Forza maggiore

12.1 Nessuna Parte sarà responsabile per ritardi o inadempimenti dovuti a eventi di Forza Maggiore.

12.2 In particolare, non matureranno Crediti di Servizio, penali, sconti o altri indennizzi a favore del Cliente in caso di indisponibilità del Servizio derivante, direttamente o indirettamente, da outage gravi e generalizzati di Microsoft Azure nell'area europea o da altri eventi straordinari fuori dal ragionevole controllo del Fornitore.

12.3 La Parte colpita da evento di Forza Maggiore dovrà darne comunicazione all'altra Parte non appena ragionevolmente possibile.

13. Riservatezza

13.1 Ciascuna Parte si impegna a mantenere strettamente riservate e a non divulgare a terzi, senza il preventivo consenso scritto dell'altra Parte, le informazioni tecniche, commerciali, organizzative, operative, economiche, contrattuali o di altra natura apprese in occasione della negoziazione, conclusione o esecuzione del presente Contratto e qualificate come riservate o che, per loro natura o per le circostanze di comunicazione, debbano ragionevolmente considerarsi riservate (le «Informazioni Riservate»).

13.2 Le Informazioni Riservate potranno essere utilizzate esclusivamente nella misura strettamente necessaria all'esecuzione del presente Contratto.

13.3 Ciascuna Parte potrà comunicare le Informazioni Riservate esclusivamente ai propri dipendenti, collaboratori, consulenti, società del gruppo, subfornitori o professionisti che abbiano effettiva necessità di conoscerle ai fini dell'esecuzione del Contratto, a condizione che tali soggetti siano vincolati da obblighi di riservatezza non meno rigorosi di quelli previsti nel presente articolo.

13.4 Non saranno considerate Informazioni Riservate quelle informazioni che la Parte ricevente dimostri: a) essere già legittimamente note prima della loro comunicazione; b) essere divenute di pubblico dominio senza violazione del presente Contratto; c) essere state legittimamente ottenute da terzi non soggetti a vincoli di riservatezza; d) essere state sviluppate autonomamente senza utilizzo delle Informazioni Riservate dell'altra Parte.

13.5 Resta salva la facoltà di ciascuna Parte di divulgare Informazioni Riservate nella misura in cui ciò sia richiesto da norme di legge, regolamenti, ordini dell'autorità giudiziaria o amministrativa competente, fermo restando, ove consentito, l'obbligo di informare preventivamente l'altra Parte.

13.6 Gli obblighi di riservatezza di cui al presente articolo resteranno efficaci per un periodo di 5 (cinque) anni dalla cessazione del Contratto, per qualsiasi causa intervenuta.

14. Protezione dei dati personali

14.1 Le Parti riconoscono che, nell'ambito dell'esecuzione del presente Contratto, potranno essere trattati dati personali ai sensi del Regolamento (UE) 2016/679 («GDPR») e della normativa nazionale applicabile in materia di protezione dei dati personali.

14.2 Il Cliente dichiara e garantisce di avere pieno titolo a caricare nel Servizio i dati, inclusi eventuali dati personali ivi contenuti, nonché di disporre di un'idonea base giuridica per il relativo trattamento.

14.3 Nella misura in cui il Fornitore tratti dati personali per conto del Cliente ai fini dell'erogazione del Servizio, il Fornitore opererà quale responsabile del trattamento ai sensi dell'art. 28 GDPR e le Parti si impegnano a sottoscrivere separato accordo di nomina a responsabile del trattamento («DPA»), che costituirà parte integrante e sostanziale del rapporto contrattuale.

14.4 In assenza di sottoscrizione del DPA, il Fornitore non sarà tenuto a svolgere trattamenti che presuppongano, ai sensi della normativa applicabile, una formale designazione a responsabile del trattamento.

14.5 Il Cliente resta in ogni caso esclusivamente responsabile: a) della liceità, correttezza, trasparenza e minimizzazione dei dati caricati; b) dell'adempimento degli obblighi informativi nei confronti degli interessati; c) dell'individuazione delle finalità e dei mezzi essenziali del trattamento effettuato mediante il Servizio; d) della gestione delle richieste degli interessati, salvo quanto di competenza del Fornitore quale responsabile del trattamento nei limiti del DPA.

14.6 Il Fornitore tratterà i dati personali esclusivamente nella misura necessaria a erogare il Servizio, garantire la sicurezza, gestire il supporto tecnico, adempiere a obblighi di legge, tutelare i propri diritti e dare esecuzione alle istruzioni documentate del Cliente, nei limiti della normativa applicabile e del DPA.

15. Limitazioni sui dati caricati

15.1 Il Cliente si impegna a non caricare nel Servizio, salvo previo accordo scritto tra le Parti e salvo quanto espressamente disciplinato nel DPA, dati rientranti nelle categorie particolari di dati personali ai sensi dell'art. 9 GDPR, ivi inclusi dati idonei a rivelare origine razziale o etnica, opinioni politiche, convinzioni religiose o filosofiche, appartenenza sindacale, dati genetici, dati biometrici trattati per identificazione univoca, dati relativi alla salute, dati relativi alla vita sessuale o all'orientamento sessuale.

15.2 Il Cliente si impegna altresì a non caricare dati personali relativi a condanne penali e reati ai sensi dell'art. 10 GDPR, salvo previo accordo scritto tra le Parti e salvo quanto espressamente consentito dalla normativa applicabile.

15.3 Resta inteso che il Fornitore non assume alcun obbligo di verifica preventiva sistematica circa la natura dei dati caricati dal Cliente e potrà, ove venga ragionevolmente a conoscenza della presenza di dati caricati in violazione del presente articolo, sospendere il trattamento interessato, richiederne la rimozione o adottare ulteriori misure ragionevoli a tutela propria, del Cliente o di terzi.

16. AI generativa e limitazioni d'uso dei dati

16.1 Il Cliente prende atto che il Servizio può includere componenti di AI generativa nell'ambito delle funzionalità di analisi, elaborazione, classificazione o strutturazione delle informazioni.

16.2 Il Fornitore non garantisce che i risultati generati mediante componenti di AI generativa siano sempre completi, accurati, esenti da errori, omissioni, ambiguità, bias o non conformità rispetto alle aspettative del Cliente, restando in ogni caso a carico del Cliente l'onere di verifica e validazione degli output prima del loro utilizzo in processi decisionali, contrattuali, legali, contabili, regolamentari o di compliance.

16.3 Il Fornitore garantisce che i dati del Cliente e i relativi output non saranno utilizzati dal Fornitore per l'addestramento, il riaddestramento o il fine-tuning di modelli generali o specifici di intelligenza artificiale.

16.4 Il Fornitore garantirà inoltre, per quanto nella propria sfera di controllo contrattuale, che eventuali terzi di cui si avvalga per l'erogazione del Servizio non utilizzino i dati del Cliente o i relativi output per finalità di training, re-training o fine-tuning di modelli di intelligenza artificiale.

16.5 Resta inteso che il Fornitore potrà trattare tali dati esclusivamente nella misura strettamente necessaria all'erogazione del Servizio, al supporto tecnico, alla sicurezza, alla continuità operativa, al rispetto di obblighi di legge e alla tutela dei propri diritti.

17. Misure tecniche e organizzative di sicurezza

17.1 Il Fornitore adotterà misure tecniche e organizzative ragionevoli e appropriate, tenuto conto dello stato dell'arte, dei costi di attuazione, della natura del Servizio e dei rischi connessi, al fine di proteggere i dati trattati contro distruzione, perdita, alterazione, divulgazione non autorizzata o accesso abusivo.

17.2 Tali misure potranno includere, a titolo esemplificativo e non esaustivo: a) controllo degli accessi logici ai sistemi e agli ambienti applicativi; b) autenticazione degli utenti e gestione profilata delle autorizzazioni; c) cifratura dei dati in transito e, ove applicabile, a riposo; d) segregazione logica degli ambienti; e) sistemi di monitoraggio, logging e tracciamento tecnico degli accessi e degli eventi rilevanti; f) procedure di patch management, vulnerability management e aggiornamento dei sistemi; g) procedure di backup, recovery e gestione degli incidenti di sicurezza; h) restrizione degli accessi al personale strettamente autorizzato.

17.3 Il Cliente prende atto che nessun sistema informatico o misura di sicurezza può garantire una protezione assoluta da eventi dolosi, vulnerabilità sconosciute, outage, malfunzionamenti o attacchi informatici e che, pertanto, il Fornitore non presta alcuna garanzia di invulnerabilità assoluta del Servizio o dell'infrastruttura sottostante.

18. Gestione degli incidenti di sicurezza e data breach

18.1 Il Fornitore si impegna a mantenere procedure ragionevoli di rilevazione, gestione e contenimento degli incidenti di sicurezza relativi al Servizio.

18.2 Qualora il Fornitore venga a conoscenza di una violazione dei dati personali trattati per conto del Cliente («Personal Data Breach»), il Fornitore ne darà comunicazione al Cliente senza ingiustificato ritardo, fornendo, nei limiti ragionevolmente disponibili al momento della comunicazione, le informazioni utili a consentire al Cliente di adempiere agli eventuali obblighi di notifica previsti dalla normativa applicabile.

18.3 Il Fornitore collaborerà ragionevolmente con il Cliente, nei limiti di quanto previsto dal DPA e della normativa applicabile, ai fini della gestione del Personal Data Breach.

18.4 Resta inteso che il Fornitore non sarà responsabile per violazioni o incidenti imputabili a comportamenti del Cliente, dei suoi Utenti Autorizzati o a sistemi, reti o ambienti non gestiti dal Fornitore.

19. Audit, verifiche e controlli

19.1 Il Cliente potrà effettuare audit on site relativi alle misure organizzative, procedurali e documentali del Fornitore rilevanti ai fini del presente Contratto e, ove applicabile, del DPA, purché: a) ne dia preavviso scritto di almeno 60 (sessanta) giorni; b) l'audit sia svolto in giorni e orari lavorativi, con modalità ragionevoli e proporzionate; c) l'audit non comprometta la sicurezza, la riservatezza o la continuità operativa del Fornitore o di altri clienti; d) il Cliente e gli eventuali auditor incaricati sottoscrivano impegni di riservatezza adeguati.

19.2 Tutti i costi e gli oneri relativi all'audit, inclusi quelli del personale interno o esterno incaricato dal Cliente, nonché gli eventuali costi ragionevoli di assistenza richiesti al Fornitore, resteranno integralmente a carico del Cliente.

19.3 Il Fornitore potrà, ove ragionevole, soddisfare in tutto o in parte le richieste del Cliente anche mediante la messa a disposizione preventiva di documentazione, questionari, attestazioni, policy o altri materiali equivalenti; resta fermo il diritto del Cliente, nei limiti del presente articolo, di richiedere audit on site.

19.4 Ogni penetration test, vulnerability assessment o verifica tecnica invasiva dei sistemi del Fornitore sarà ammessa esclusivamente previo separato accordo scritto tra le Parti.

20. Uso consentito del Servizio e divieti

20.1 Il Cliente si impegna a utilizzare il Servizio esclusivamente per finalità lecite, connesse alla propria attività aziendale e in conformità al presente Contratto, alla normativa applicabile e alla documentazione d'uso eventualmente fornita dal Fornitore.

20.2 È fatto divieto al Cliente, direttamente o indirettamente, di: a) copiare, modificare, adattare, tradurre, decompilare, disassemblare o sottoporre a reverse engineering il Servizio o parti di esso, salvo quanto inderogabilmente consentito dalla legge; b) aggirare, compromettere o tentare di compromettere i meccanismi di sicurezza, autenticazione, segregazione logica, controllo accessi o limitazione d'uso del Servizio; c) effettuare test di carico, stress test, vulnerability assessment o penetration test sulla Piattaforma senza autorizzazione scritta del Fornitore; d) utilizzare il Servizio per finalità illecite o in violazione della normativa applicabile; e) caricare dati in violazione dell'articolo 15.

20.3 Il Fornitore potrà sospendere l'accesso al Servizio in caso di violazione del presente articolo, fermo restando ogni ulteriore diritto o rimedio previsto dal Contratto o dalla legge.

21. Proprietà intellettuale

21.1 Tutti i diritti di proprietà intellettuale relativi a {p['name']}, alla Piattaforma, al software, agli algoritmi, ai template standard, alla documentazione e ai relativi aggiornamenti restano di esclusiva titolarità del Fornitore o dei suoi licenzianti.

21.2 Il Cliente conserva i diritti sui propri dati e contenuti caricati, fermo restando che il presente Contratto non comporta alcun trasferimento di diritti sul software.

22. Limitazione di responsabilità

22.1 Il Fornitore non sarà responsabile per danni indiretti, perdita di profitto, perdita di chance, perdita di dati, fermo attività o danni reputazionali, salvo il dolo o la colpa grave.

22.2 Il Fornitore non sarà altresì responsabile per errori o inesattezze dei risultati derivanti da: a) qualità insufficiente dei dati di input; b) configurazioni errate o incomplete; c) uso improprio del Servizio.

22.3 Salvo il dolo o la colpa grave e fatti salvi i limiti inderogabili di legge, la responsabilità complessiva del Fornitore, a qualunque titolo insorta in relazione al presente Contratto, non potrà eccedere un importo pari al {p['resp_limit_pct']}% del totale dei Canoni di Licenza corrisposti dal Cliente nel corso dell'anno solare precedente a quello in cui si è verificato l'evento dannoso.

22.4 Qualora l'evento dannoso si verifichi nel corso del primo anno di efficacia del Contratto, ovvero in assenza di un intero anno solare precedente di riferimento, il limite massimo di responsabilità sarà determinato in misura proporzionale, assumendo come base di calcolo il {p['resp_limit_pct']}% della quota parte dei Canoni di Licenza fissi maturati e corrisposti dal Cliente sino alla data dell'evento dannoso.

22.5 I Crediti di Servizio eventualmente riconosciuti ai sensi degli articoli 6, 7 e 8 costituiscono l'unico rimedio economico del Cliente in relazione al mancato rispetto dei relativi SLA.

23. Durata

23.1 Il Contratto ha durata iniziale di 12 mesi dalla data di efficacia.

23.2 Alla scadenza, esso si rinnoverà automaticamente per successivi periodi di pari durata, salvo disdetta comunicata con preavviso scritto di almeno 30 giorni.

24. DPA

24.1 Le Parti convengono che gli aspetti di dettaglio relativi al trattamento dei dati personali per conto del Cliente, inclusi eventuali sub-responsabili, istruzioni documentate, misure tecniche e organizzative, assistenza all'esercizio dei diritti, gestione dei breach e trasferimenti internazionali, saranno disciplinati in un separato Data Processing Agreement ai sensi dell'art. 28 GDPR.

25. Legge applicabile e foro competente

25.1 Il presente Contratto è regolato dalla legge italiana.

25.2 Per ogni controversia sarà competente in via esclusiva il Foro di Milano.

Milano, {p['date']}

Per il Fornitore
______________________________

Per il Cliente
______________________________"""


def create_docx(text, filepath):
    """Create a .docx file from contract text."""
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    lines = text.split("\n")
    for i, line in enumerate(lines):
        stripped = line.strip()
        if not stripped:
            continue

        para = doc.add_paragraph()

        # Title (first line)
        if i == 0:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(stripped)
            run.bold = True
            run.font.size = Pt(14)
        # Product name (second non-empty line)
        elif i == 1:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(stripped)
            run.bold = True
            run.font.size = Pt(13)
        # Section headers (numbered sections like "1. Premesse")
        elif stripped and len(stripped) < 80 and stripped[0].isdigit() and ". " in stripped[:4] and not stripped[0:2].count(".") > 0:
            run = para.add_run(stripped)
            run.bold = True
            run.font.size = Pt(12)
        # Signature lines
        elif stripped.startswith("______"):
            para.add_run(stripped)
        elif stripped.startswith("Milano,"):
            para.add_run("")  # spacing
            para = doc.add_paragraph()
            para.add_run(stripped)
        else:
            para.add_run(stripped)

    doc.save(filepath)


def main():
    total = 0
    for product in PRODUCTS:
        folder_path = os.path.join(BASE_DIR, product["folder"])
        os.makedirs(folder_path, exist_ok=True)

        for client in product["clients"]:
            filename = f"contratto_{client}.docx"
            filepath = os.path.join(folder_path, filename)
            text = build_contract_text(product, client)
            create_docx(text, filepath)
            total += 1

        print(f"  OK {product['folder']}: {len(product['clients'])} contratti ({product['name']} - {product['category']})")

    print(f"\nTotale: {total} file generati.")


if __name__ == "__main__":
    main()
